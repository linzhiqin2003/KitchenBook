"""Comprehensive concept quiz: questions first, answers at the end."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.line_spacing = 1.15
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

qn = [0]
answers = []  # collect all answers

def Q(text, marks, mock_label=None):
    qn[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    prefix = f"Q{qn[0]}. "
    if mock_label:
        prefix = f"Q{qn[0]}. [MOCK {mock_label}] "
    r = p.add_run(prefix)
    r.bold = True
    p.add_run(text)
    r2 = p.add_run(f"  [{marks} marks]")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return qn[0]

def sub(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(text)

def A(num, bullets, total_marks):
    answers.append((num, bullets, total_marks))

# ═══ Cover ═══
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("COMSM0166 — Software Engineering"); r.font.size = Pt(14); r.bold = True
doc.add_paragraph()
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Comprehensive Concept Quiz — Chapters 1–5"); r2.font.size = Pt(18); r2.bold = True; r2.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
doc.add_paragraph()
for line in ["Total: ~150 marks · 65+ questions","Format: concept recall + mock exam originals","","Questions are in the front section.","Answer key is at the end — cover it while testing yourself."]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(line)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART 1: L1 INTRODUCTION
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 1 — Introduction to Software Engineering", level=1)

n=Q("List three reasons behind the 'Software Crisis'.", 3)
A(n,["Software is getting larger and larger. [1]","Software is getting more complex (complex domains like human behaviour, systems dependency like energy and cars). [1]","Time to market is shorter than ever. [1]"],3)

n=Q("List six reasons why software projects fail.", 3)
A(n,["Poor requirements / over-ambitious requirements / unnecessary requirements. [0.5 each, max 1.5]","Over budget. [0.5]","Contract management issues. [0.5]","End-user training gaps / operational management failures. [0.5]","Note: 'bad developers' is NOT the main problem."],3)

n=Q("Define Software Engineering.", 2)
A(n,["Engineering: cost-effective solutions to practical problems by applying scientific knowledge to building things for people. [1]","Software Engineering deals with the design, development, testing, and maintenance of software applications. Engineers apply engineering principles and programming knowledge to build software solutions for end users. [1]"],2)

n=Q("List five collaboration tools and techniques mentioned in the course.", 2)
A(n,["Modelling and diagramming (to reach consensus on design). [0.5]","GitHub (sharing documents for effective collaboration). [0.5]","Kanban (task allocation and progress monitoring). [0.5]","Test-driven development (stop others breaking your code). [0.5]","Other techniques: feel free to experiment and explore."],2)

n=Q("List the seven Software Development Tasks.", 3)
A(n,["Requirements Analysis [0.5]","Planning [0.5]","Design (high level and detailed) [0.5]","Development [0.5]","Testing [0.5]","Deployment [0.5]","Operation and Maintenance [0.5]","(These tasks are combined in various sequences, making up different SDLCs.)"],3)

n=Q("Name four examples of Software Development Life Cycles.", 2)
A(n,["Waterfall [0.5]","Agile [0.5]","V-Model [0.5]","Spiral [0.5]"],2)

n=Q("List and briefly describe the five stages of the Waterfall SDLC.", 5)
A(n,["Requirements Definition: gathering functional (what the system should do) and non-functional (quality properties) requirements. [1]","System and Software Design: deciding how parts interact, dividing work between team members. [1]","Implementation and Unit Testing: parallel working, code sharing, API partitioning, versioning, automated testing, docs. [1]","Integration and System Testing: combining modules, verifying the whole system works together. [1]","Operation and Maintenance: deploying, ongoing support, bug fixes, updates. [1]"],5)

n=Q("Define functional requirements and non-functional requirements (in the context of Waterfall).", 2)
A(n,["Functional: what the system should do (services, reactions to inputs, behaviours in situations, also what NOT to do). [1]","Non-functional: quality properties of system operation, e.g., security, ease of use, response time. [1]"],2)

n=Q("Explain the difference between Verification and Validation. Give an example where software passes verification but fails validation.", 5)
A(n,["Verification: 'Are you building it right?' — checks software complies with specifications, constraints, and regulations. Demonstrates system meets specifications. [2]","Validation: 'Are you building the right thing?' — checks software meets actual needs of customers/stakeholders. Demonstrates system meets user needs. [2]","Example: a system built exactly per specification (passes verification) but the specification did not address user needs (fails validation). Can pass verification but fail validation if built as per specs yet specs don't address user needs. [1]"],5)

n=Q("List three advantages and three disadvantages of the Waterfall model.", 3)
A(n,["Advantages (0.5 each): Simple to use and understand; every phase has defined result and process review; development stages go one by one; perfect for projects with clear agreed requirements; easy to determine key points; easy to classify and prioritize tasks.","Disadvantages (0.5 each): Software ready only after last stage; high risks and uncertainty; misses complexity due to interdependence; not suited for long-term projects with changing requirements; progress hard to measure during development; integration done at very end — can't identify problems in advance."],3)

n=Q("Define the Software Development Lifecycle (SDLC).", 2)
A(n,["A structured and iterative methodology used by development teams to build, deliver and maintain high-quality and cost-effective software systems. [2]"],2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART 2: L2 AGILE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 2 — Agile Software Development", level=1)

n=Q("What is Agile Software Development?", 1)
A(n,["Agile is a way of thinking about software development. [1]"],1)

n=Q("State the four key values of the Agile Manifesto.", 4)
A(n,["Individuals and interactions over processes and tools. [1]","Working software over comprehensive documentation. [1]","Customer collaboration over contract negotiation. [1]","Responding to change over following a plan. [1]"],4)

n=Q("List four things coders LIKE and four things coders DISLIKE, as described in the Agile philosophy.", 4)
A(n,["Like: Writing quality code [0.5]; Ticking things off their to-do list [0.5]; Impressing clients by showing working software [0.5]; (any of the above) [0.5]","Dislike: Writing extensive documentation [0.5]; Committing to a final design in advance [0.5]; Being micromanaged [0.5]; Working to big, immovable deadlines [0.5]"],4)

n=Q("List six Agile principles that satisfy clients' needs and six that satisfy coders' needs.", 6)
A(n,["Client-side (1 mark per pair, 0.5 each): Highest priority = satisfying client (deliver working software early and continuously); Embrace change (even late in cycle); Collaborate every day with client; Use face-to-face communication; Deliver working software frequently; Minimise unnecessary work.","Coder-side: Work at steady sustainable pace; Rely on self-organising teams; Teams reflect regularly on performance; Progress measured by working code produced; Continuous attention to technical excellence; Build teams around motivated individuals."],6)

# MOCK Q1
n=Q("Describe the ethos underlying Extreme Programming. [5 marks]\nBriefly describe five practices that are used in Extreme Programming. [5 marks]", 10, "Q1")
A(n,["Ethos (1 mark each): Simple design (simplest way to implement features); Sustainable pace (effort is constant and manageable); Coding standards (teams follow agreed style and format); Collective ownership (everyone owns all the code); Whole team approach (everyone is included in everything).","Practices (1 mark each): Pair programming (two heads are better than one); Test driven (ensure code runs correctly); Small releases (deliver frequently, get feedback); Continuous integration (ensure system is operational); Refactoring (restructure system when things get messy)."],10)

n=Q("Describe pair programming in detail: the two roles, key rules, and benefits.", 4)
A(n,["Helm: uses keyboard and mouse, does the coding. [1]","Tactician: thinks about implications and potential problems, positioned to recommend refactoring. [1]","Communication is essential; the pair doesn't 'own' the code — anyone can change it. [0.5]","Pairings can and should evolve at any time. All code is reviewed as it is written. [0.5]","Pair programming facilitates project communication. [0.5]"],4)

n=Q("Explain Test-Driven Development: the core rules and three benefits.", 3)
A(n,["Rules: Tests are written before any code and drive all development; a programmer's job is to write code to pass tests; if there's no test for a feature, it is not implemented; tests are informed by requirements. [1.5]","Benefits: Code coverage — all code has at least one test. [0.5]; Simplified debugging — failing test must be caused by last change. [0.5]; System documentation — tests describe what code should do. [0.5]"],3)

n=Q("List and define the five key Scrum concepts.", 3)
A(n,["The Scrum: a stand-up daily meeting of the entire team. [0.5]","Scrum Master: team leader, facilitates agile processes. [0.5]","Sprint: a short, rapid development iteration. [0.5]","Product Backlog: to-do list of jobs that need doing. [0.5]","Product Owner: the client (or their representative). [0.5]"],3)

n=Q("Describe the Kanban board: what it is, how it works, and what columns it might have.", 3)
A(n,["A flexible 'to do' list tool (originally post-it notes on whiteboard, now digital tools like Jira). [1]","Issues progress through various states from 'To do' to 'Done'. [1]","Columns might include: Backlog, Being Verified, Awaiting Integration, Done. Must always have a 'Done' column. Do what works for your team. [1]"],3)

n=Q("List four problems or limitations of Agile development.", 4)
A(n,["Hard to draw up legally binding contracts — full specification never written in advance. [1]","Good for greenfield development but not effective for brownfield (legacy systems). [1]","Works well for small co-located teams, but challenging for large distributed development. [1]","Relies on knowledge of developers — what if they aren't around (holidays, illness, turnover)? [1]"],4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART 3: L3 REQUIREMENTS ENGINEERING
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 3 — Requirements Engineering", level=1)

n=Q("What are system requirements? Distinguish functional and non-functional requirements with examples.", 3)
A(n,["System requirements specify a system in terms of user observation, not implementation. They record descriptions of features and constraints. [1]","Functional: specify user interactions — what the system does (services, reactions to inputs, behaviour in situations, what it should NOT do). [1]","Non-functional: specify other system properties — constraints on services, often apply to whole system (security, usability, performance, etc.). [1]"],3)

n=Q("List five general problems that make requirements engineering necessary.", 3)
A(n,["Inconsistent terminology: people express needs in their own words. [0.5]","Conflicting needs for the same system. [0.5]","People frequently don't know what they want (or can't explain). [0.5]","Requirements change quite frequently. [0.5]","Relevant people/information may not be accessible. [0.5]","(Requirements serve as: communication mechanism, instructions, and acceptance criteria.) [0.5]"],3)

n=Q("What three properties must good requirements have?", 2)
A(n,["Unambiguous / Precise. [0.5]","Sufficiently accessible / Measurable. [0.5]","Understandable / Clear. [0.5]","(To fairly assess whether the team produced something matching what was asked for.) [0.5]"],2)

n=Q("What is the Onion Model? List at least six stakeholder roles it identifies.", 3)
A(n,["The Onion Model is a template for identifying stakeholders, organized in concentric layers from the product outward. [1]","Roles include: Normal Operator, Functional Beneficiary, Champion, Sponsor, Purchaser, Developer, Maintenance Operator, Operational Support, Interfacing System Owner, Negative Stakeholder, Regulator, Politician, The Public. [2 — 0.25 each, up to 2]"],3)

n=Q("List seven methods for stakeholder identification.", 2)
A(n,["Clients; Documentation (e.g. organisation charts); Templates (e.g. onion model); Similar projects; Analysing context of the project; Surrogate stakeholders (legal, unavailable, mass users); Negative stakeholders. [~0.3 each]"],2)

n=Q("Define an Epic and list its characteristics.", 2)
A(n,["An Epic is a high-level requirement (no detail), focused on business or user value, large and/or complex at scope, needs to be broken down into smaller pieces to implement. [2]"],2)

n=Q("Write the User Story template and give one example.", 2)
A(n,["Template: As a <type of user>, I want to <some goal> so that <some reason>. [1]","Example: As a user with low vision, I want to resize text so that I can read comfortably. [1]"],2)

# MOCK Q2
n=Q("2.a) How is INVEST used for agile requirements [1 mark] and what does each letter in INVEST mean [6 marks]?\n2.b) Which three INVEST criteria does this user story fail: \"As a disabled user, I want to update the app settings so that I can use the app comfortably\". [3 marks]", 10, "Q2")
A(n,["2.a) INVEST is used to evaluate the quality of user stories (how well a user story is written). [1]","I — Independent: Can be worked on separately from other stories. [1]","N — Negotiable: Flexible and open to discussion. [1]","V — Valuable: Delivers clear value to the user. [1]","E — Estimable: Can be estimated for effort. [1]","S — Small: Small enough to complete within a sprint. [1]","T — Testable: Has clear criteria to determine if it's done. [1]","2.b) Fails Small: more of an initiative/epic — many disability types need multiple sprints. [1]; Fails Estimable: not clear how to estimate, depends on which disabilities. [1]; Fails Testable: not clear what to test, depends on which disabilities for acceptance tests. [1]"],10)

n=Q("Write the Acceptance Criteria template and list the four qualities good criteria must have.", 2)
A(n,["Template: Given <initial context/precondition>, When <action/event>, Then <expected outcome>. [1]","Qualities: Clear (unambiguous), Testable (can verify), Measurable (quantitative/qualitative), Atomic (independent, checked by itself). [1]"],2)

n=Q("Explain the hierarchy: Initiative → Epic → User Story. Compare them in terms of scope and timeframe.", 3)
A(n,["Initiative: strategic objective, important business outcome, spans multiple epics and teams/departments. [1]","Epic: large strategic goal, spans multiple sprints, needs breakdown into smaller pieces. [1]","User Story: specific feature/functionality, completed within a single sprint. [1]"],3)

n=Q("Explain the MoSCoW prioritization method.", 2)
A(n,["Must-Have: essential. [0.5]","Should-Have: important. [0.5]","Could-Have: nice to have. [0.5]","Won't-Have: out of scope at present. [0.5]"],2)

n=Q("Explain the Value vs Effort prioritization matrix.", 2)
A(n,["High value, Low effort → do first. [0.5]","High value, High effort → do next. [0.5]","Low value, Low effort → do if there is time. [0.5]","Low value, High effort → avoid. [0.5]"],2)

n=Q("List seven requirements elicitation techniques.", 2)
A(n,["Interviews, Observations, Surveys, Current documentation, Similar products and solutions, Co-design, Prototyping. [~0.3 each]"],2)

n=Q("What is a Use Case Model? What two parts does it consist of?", 3)
A(n,["A use-case model describes functional requirements in terms of use cases, links stakeholder needs to software requirements, and serves as a planning tool. [1.5]","Consists of: Use-case diagrams (visual representation) and Use-case specifications (text representation). [1.5]"],3)

n=Q("What is an Actor? List five properties of actors.", 3)
A(n,["An actor represents anything that interacts with the system. [1]","Can represent a human, machine, or another system. [0.5]","Can actively interchange information with the system. [0.5]","Can be a giver of information or a passive recipient. [0.5]","Actors are NOT part of the system — they are EXTERNAL. [0.5]"],3)

n=Q("What is a Use Case? What does it describe?", 2)
A(n,["A use case defines a set of use-case instances, where each instance is a sequence of actions the system performs that yields an observable result of value to a particular actor. [1]","It models a dialogue between one or more actors and the system, describing actions the system takes to deliver value. [1]"],2)

n=Q("What is a use-case specification? What does it contain?", 2)
A(n,["A requirements document containing the text of a use case. [0.5]","Contains: description of flow of events (actor-system interaction), preconditions, postconditions, special requirements, key scenarios, subflows. [1.5]"],2)

n=Q("Explain basic flow vs alternative flows in use cases.", 2)
A(n,["Basic flow: the successful scenario from start to finish. One per use case. [1]","Alternative flows: regular variants, odd cases, exceptional/error flows. Many per use case. [1]"],2)

n=Q("List the three checkpoints for validating use cases.", 2)
A(n,["Each use case is independent of the others. [0.5]","No use cases have very similar behaviours or flows. [0.5]","No part of the flow has already been modelled as another use case. [0.5]"],2)

n=Q("List three benefits of a Use-Case Model.", 1)
A(n,["Communication, Identification, Testing. [~0.33 each]"],1)

# MOCK Q5
n=Q("You are employed by WeCode Ltd to work with NiceHair, a hairdresser's business. Your team develops a software system for tracking hairstyling products.\n5.1 Write 2 epics and 3 user stories per epic for this system. [16 marks]\n5.2 Write acceptance criteria for each user story. [6 marks]\n5.3 What are the key roles within this agile team? [3 marks]", 25, "Q5")
A(n,["5.1 Epic 1: Product Catalog Management — (1) As a manager, I want to add new products to the catalog so that customers and staff can see available offerings. (2) As a hairstylist, I want to search products by name/category/brand so I can quickly find the right product. (3) As an owner, I want to update product details (pricing, availability) so the catalog stays accurate. [8 marks]","5.1 Epic 2: Inventory and Stock Management — (1) As a manager, I want to set minimum stock levels so I receive alerts when stock is low. (2) As a hairstylist, I want to check product availability before recommending it. (3) As an owner, I want to generate stock reports to track usage and reorder efficiently. [8 marks]","5.2 Acceptance Criteria (Given/When/Then for each story, 1 mark each). E.g., 'Given the manager is logged in, When they enter product details and click Save, Then the new product appears in the catalog.' [6 marks]","5.3 Product Owner (sets vision, prioritises backlog) [1]; Scrum Master (facilitates processes, removes obstacles) [1]; Developers (write, test, deploy code) [1]. Also acceptable: testers, UX designer, business analyst."],25)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART 4: L4 OOD
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 4 — Object-Oriented Design", level=1)

n=Q("What is Object-Oriented Software? Distinguish Classes from Objects.", 2)
A(n,["OO Software: structure software around objects and their interactions. Objects (real-world entities) contain state and behaviour. [1]","Classes: blueprints for types of objects (e.g., Car). Objects: individual instances (e.g., Ford Fusion with registration AB25CDF). [1]"],2)

n=Q("List five reasons why we do design.", 2)
A(n,["Organise ideas; Plan work; Build understanding of system structure and behaviour; Communicate with development team; Help future maintenance team understand. [0.4 each]"],2)

n=Q("Define Encapsulation and explain why it matters. Give a code example.", 3)
A(n,["Encapsulation: encloses the data and behaviour of objects within that object. [1]","Prevents unauthorised access. Controlled access: object can be changed only by permitted methods. Protects data integrity. [1]","Example: class Animal { private String name; public String getName() { return name; } } — name is private, accessed only via getName(). [1]"],3)

n=Q("Define Abstraction and explain its benefit.", 2)
A(n,["Abstraction: focus on core concerns, leave out unnecessary detail, expose only essential information, hide complexity. [1]","Benefit: promotes loose coupling — objects interact via abstract interfaces, not concrete implementations. [1]"],2)

n=Q("Define Inheritance. State two benefits.", 2)
A(n,["Inheritance: inherit properties and behaviour from another class ('is-a' relationship). [1]","Benefits: reuse code (reduce re-writing), reduce errors and inconsistency. [1]"],2)

n=Q("Define Polymorphism. Give a code example.", 2)
A(n,["Polymorphism: allows substitution of subclass for superclass. Behaviour belongs to the subclass (dynamic method dispatch). Enables extensibility, flexible/modular systems. [1]","Example: Animal animal1 = new Dog(); animal1.makeSound(); // Output: Woof! — the subclass method is called. [1]"],2)

n=Q("Define Composition. Distinguish it from Aggregation.", 3)
A(n,["Composition: build an object from other objects. Enables reuse and incremental construction. Non-shared aggregation with strong ownership and coincident lifetimes — parts cannot survive the whole. [1.5]","Aggregation: a whole-part ('is a part-of') relationship where parts CAN exist independently. Can be shared (part belongs to multiple wholes) or non-shared. [1.5]"],3)

n=Q("What does a UML Class Diagram represent? What does a class contain?", 2)
A(n,["A class diagram shows a static view of a system. [0.5]","A class contains: Class Name, Attributes (visibility, name, type), Methods (visibility, name, parameters, return type). [1.5]"],2)

n=Q("List the four visibility/access modifiers in UML with their meanings.", 2)
A(n,["+ public: available to all. [0.5]","- private: available only within the class. [0.5]","# protected: available to subclasses and same-package classes. [0.5]","~ package: available only within the package. [0.5]"],2)

n=Q("Write the notation for attributes and operations in UML.", 2)
A(n,["Attributes: [visibility] name [: type] [multiplicity] [= value] [{property}]. Static attributes are underlined. [1]","Operations: [visibility] name ([parameter-list]) [: return-type] [{property}]. Static operations are underlined. Parameters: direction name : type [multiplicity] = value [{property}]. [1]"],2)

n=Q("How do you find classes? Describe the Grammatical Parse approach.", 2)
A(n,["Identify nouns from existing text (descriptions, requirements). [1]","Narrow down by removing: duplicates and variations (synonyms), irrelevant items, out-of-scope items. [1]"],2)

n=Q("Define Association and Multiplicity in class diagrams.", 2)
A(n,["Association: a semantic/structural relationship between two or more classifiers, specifying connections among their instances. [1]","Multiplicity: the number of instances one class relates to ONE instance of another class. Two multiplicity decisions per association (one at each end). [1]"],2)

n=Q("List all multiplicity notations and their meanings.", 3)
A(n,["1 = exactly one. [0.5]","0..* or * = zero or more. [0.5]","1..* = one or more. [0.5]","0..1 = zero or one (optional). [0.5]","2..4 = specified range. [0.5]","2, 4..6 = multiple disjoint ranges. [0.5]"],3)

n=Q("Explain the difference between Generalization and Aggregation.", 2)
A(n,["Generalization: 'is a kind of' — subclass inherits from superclass, shares properties/behaviour. [1]","Aggregation: 'part of' — whole-part relationship, one object contains others. [1]"],2)

n=Q("What is the difference between Abstract and Concrete classes?", 2)
A(n,["Abstract classes cannot have any direct objects (instances). They serve as blueprints. [1]","Concrete classes can have objects — all actual instances are of concrete classes. [1]"],2)

n=Q("Define Navigability in class diagrams.", 1)
A(n,["Indicates that it is possible to navigate from an associating class to the target class using the association. [1]"],1)

n=Q("What is the purpose of a Sequence Diagram? List the key elements.", 3)
A(n,["Purpose: models how objects collaborate and interact over time through messages. Good for real-time specifications and complex scenarios. [1]","Elements: Participants (actors and objects); Lifeline (vertical dashed line = existence over time); Vertical time axis (flow of time top to bottom); Messages (horizontal arrows = requests, may contain parameters); Return messages (dashed arrows, may contain return value or void). Sync vs async arrowheads. [2]"],3)

n=Q("What interaction frames can be used in Sequence Diagrams? What can SDs model?", 2)
A(n,["Interaction frames: alt (branching/conditional), loop (iteration). [1]","SDs can model: simple sequential flow, branching, iteration, recursion, concurrency. May specify primary, variant, and exception scenarios. [1]"],2)

# MOCK Q6
n=Q("Design an ATM system for withdrawals, deposits, and balance inquiries. ATM authenticates via card number and PIN. Transactions include ID, date, type, amount, post-transaction balance.\n6.1 Sketch the class diagram with main classes and relationships. [5 marks]\n6.2 List 2 key attributes and 2 key methods for an account. [4 marks]\n6.3 Explain access modifiers for these attributes and methods, and why choices matter. [2 marks]\n6.4 Demonstrate 2 multiplicity examples and explain them. [4 marks]\n6.5 Provide examples and one-sentence explanations for: encapsulation, abstraction, inheritance, polymorphism, aggregation/composition. [10 marks]", 25, "Q6")
A(n,["6.1 Main classes (1 mark each, need ≥3): ATM, Customer, Bank, Account, ATM Transactions, Current Account, Savings Account. Show attributes [1], methods [1], relationships [1], multiplicities [1].","6.2 Attributes: number (account number) [1], balance [1]. Methods: deposit() [1], withdraw() [1].","6.3 Attributes (number, balance): private (-) — enforce encapsulation, prevent direct access, use getters/setters for security. [1] Methods (deposit, withdraw): public (+) — need to be accessible to other classes (ATM, Customer), validation before modifying balance. [1]","6.4 Customer → Account (1 to 1,2): a customer can have one or two accounts, each belongs to one customer. [2] Account → ATM Transactions (1 to *): each account can have multiple transactions. [2]","6.5 Encapsulation: Account has private balance/number with public deposit()/withdraw() — protects sensitive data. [2] Abstraction: Account could be abstract, allowing other types without replicating common logic. [2] Inheritance: Current Account and Savings Account inherit from Account — reuse common features, add specialised behaviour. [2] Polymorphism: withdraw() behaves differently in Current (allows overdraft) vs Savings (restricts) — single method, different implementations. [2] Aggregation/Composition: Bank is an aggregation of ATM, Account, Customer — bigger entities built from parts. [2]"],25)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# PART 5: L5 PROJECT MANAGEMENT
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 5 — Project Management", level=1)

n=Q("Describe the key practices in Agile Management, organized by the three main roles.", 4)
A(n,["Iterative and incremental delivery via sprint cycles (2-4 weeks). [1]","Product Owner: Epics, User Stories, Backlog Grooming. [1]","Development Team: Continuous Integration/Delivery, Pair Programming, Test-Driven Development. [1]","Scrum Master: Daily stand-ups, sprint planning, sprint reviews, retrospectives. Also: Kanban (limiting WiP, transparency of progress). [1]"],4)

n=Q("Why is measurement central to quality? What is 'measurement'?", 3)
A(n,["Measurement helps: plan project time and effort (for team and customer), identify which parts need more testing, evaluate productivity. [1.5]","Measurement: attributing values to objects (e.g., fuel efficiency, goals scored). Can use values as basis for comparison. Can use comparisons to make better decisions. [1.5]"],3)

n=Q("Why is measurement difficult in Software Engineering?", 1)
A(n,["Most entities are difficult to measure reliably. Difficult or impossible to 'pin down' a single value. [1]"],1)

n=Q("List the eight top-level quality characteristics in the ISO/IEC 25010 quality model.", 4)
A(n,["Functional Suitability (completeness, correctness, appropriateness). [0.5]","Performance Efficiency (time behaviour, resource utilisation, capacity). [0.5]","Compatibility (co-existence, interoperability). [0.5]","Usability (learnability, operability, accessibility, etc.). [0.5]","Security (confidentiality, integrity, non-repudiation, authenticity, accountability). [0.5]","Maintainability (modularity, reusability, analysability, modifiability, testability). [0.5]","Reliability (maturity, availability, fault tolerance, recoverability). [0.5]","Portability (adaptability, installability, replaceability). [0.5]"],4)

n=Q("Compare Lines of Code (LoC) and Cyclomatic Complexity as white-box metrics.", 4)
A(n,["LoC: Easy to compute and understand, widely used, often sufficient for approximate size. [1]","LoC limitations: ignores logical/architectural complexity, highly language-specific, not all lines equal, comments/blank lines ambiguity. [1]","Cyclomatic Complexity: V(G) = E − N + 2P (E=edges, N=nodes, P=procedures). Measures number of independent paths through code. [1]","An independent path introduces at least one new statement/condition. Better than LoC for measuring logical complexity. [1]"],4)

n=Q("Explain Story Points and how Planning Poker works.", 4)
A(n,["Story Points: an informal agile unit of size measurement, usually 1-10. Derived collectively by the whole team at sprint planning meetings. Based on 'Wisdom of the Crowds' — collective estimates are better than individual. [2]","Planning Poker: each member gets numbered cards following Fibonacci (1,3,5,8,13,20...). Larger tasks harder to estimate precisely. Each member also gets a '?' card. All reveal simultaneously, discuss divergences, repeat up to 3 iterations. [2]"],4)

n=Q("Define Team Velocity and explain how it is used.", 2)
A(n,["Team Velocity: number of estimated story points implemented per sprint. Can be derived from previous sprints (e.g., average of last x sprints). [1]","Used to estimate: time required to complete project; target number of stories completable in a sprint. [1]"],2)

n=Q("What do Burn-Down Charts show? How do feedback loops work?", 2)
A(n,["Burn charts plot remaining work (story points) vs time (sprints), showing estimated vs actual lines. [1]","Feedback loops: after each sprint, compare estimated vs actual progress. Discrepancies are discussed in retrospectives to drive improvement. [1]"],2)

# MOCK Q3
n=Q("Describe the key characteristics (purpose, who attends, outcome, when it happens, what is discussed) of Sprint Retrospective.", 10, "Q3")
A(n,["Purpose: To inspect the process, reflect on team performance, and identify improvements. [2]","Attendees: Development team and Scrum Master (Product Owner may join but is not required). [2]","Outcome: Actionable items aimed at improving team processes in future sprints. [2]","When It Happens: At the end of the sprint, immediately after or separate from the sprint review. [2]","What Is Discussed: What went well during the sprint; what didn't go well; team challenges and blockers; how to improve collaboration and productivity. [2]"],10)

n=Q("Describe the key characteristics of Sprint Review: Purpose, Focus, Attendees, Output, Key Question.", 5)
A(n,["Purpose: To inspect the product increment and gather feedback from stakeholders. [1]","Focus: The product and deliverables completed during the sprint. [1]","Attendees: Dev team, Product Owner, Scrum Master, stakeholders, customers, users. [1]","Output: Feedback on product, potential changes to product backlog or requirements. [1]","Key Question: 'Does the product meet stakeholder expectations, and what adjustments are needed?' [1]"],5)

n=Q("List the four categories of the Scrum framework: Roles, Ceremonies, Artifacts, Agreements. Name all items in each.", 4)
A(n,["Roles: Product Owner, Scrum Master, Team. [1]","Ceremonies: Sprint Planning, Daily Scrum, Sprint Review, Sprint Retrospective. [1]","Artifacts: Product Backlog, Sprint Backlog, Working Software. [1]","Agreements: Working Agreement, Definition of Ready, Definition of Done. [1]"],4)

n=Q("Compare Sprint Review and Sprint Retrospective across: Purpose, Attendees, Focus, Output, Key Question.", 5)
A(n,["Review — Purpose: inspect product increment + gather feedback. Retro — Purpose: inspect process + identify improvements. [1]","Review — Attendees: dev team, PO, SM, stakeholders, customers. Retro — Attendees: dev team + SM (PO optional). [1]","Review — Focus: the product. Retro — Focus: process, collaboration, workflow. [1]","Review — Output: feedback on product, backlog changes. Retro — Output: actionable items for process improvement. [1]","Review — Key Q: 'Does product meet expectations?' Retro — Key Q: 'How can we improve how we work?' [1]"],5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# ANSWER KEY
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("ANSWER KEY", level=1)
p = doc.add_paragraph()
r = p.add_run("Tear off or cover this section while testing yourself.")
r.italic = True
r.font.color.rgb = RGBColor(0x88,0x88,0x88)
doc.add_paragraph()

total_marks = 0
for (num, bullets, marks) in answers:
    total_marks += marks
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"Q{num}")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00,0x5A,0x9E)
    r2 = p.add_run(f"  [{marks} marks]")
    r2.font.color.rgb = RGBColor(0x99,0x99,0x99)
    r2.font.size = Pt(10)

    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.text = ""
        bp.add_run(b).font.size = Pt(10)

    # thin separator
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(2)
    rr = sp.add_run("─" * 50)
    rr.font.color.rgb = RGBColor(0xDD,0xDD,0xDD)
    rr.font.size = Pt(7)

# End
doc.add_page_break()
doc.add_heading("End of Quiz", level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Total: {qn[0]} questions · {total_marks} marks")
r.bold = True
r.font.size = Pt(14)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SEDP_Concept_Quiz_v2_L1-L5.docx")
doc.save(out)
print(f"Generated {qn[0]} questions, {total_marks} marks → {out}")
