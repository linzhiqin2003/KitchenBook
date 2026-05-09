"""Generate a concept-recall quiz Word document for SEDP L1-L5."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    doc.styles[f"Heading {level}"].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

q_num = [0]

def q(text, marks):
    q_num[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f"Q{q_num[0]}. ")
    r.bold = True
    p.add_run(text)
    r2 = p.add_run(f"  [{marks} marks]")
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return q_num[0]

def ans_header():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Answer:")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)

def ans(text, indent=0):
    from docx.shared import Inches
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Inches(0.3 + indent * 0.3)
    p.text = ""
    p.add_run(text)

def ans_text(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic

def sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run("─" * 55)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r.font.size = Pt(7)

# ═══ Cover ═══
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("COMSM0166 — Software Engineering")
r.font.size = Pt(14)
r.bold = True

doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Concept Recall Quiz — Chapters 1–5")
r2.font.size = Pt(18)
r2.bold = True
r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()
for line in [
    "Total: 100 marks",
    "Format: Short-answer concept definitions",
    "Purpose: Test memorization of key terms and frameworks",
    "",
    "Tip: Cover the answers and test yourself. Check one concept at a time.",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    if "100" in line:
        run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# L1: INTRODUCTION (15 marks)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 1 — Introduction to Software Engineering  (15 marks)", level=1)

q("Define Software Engineering in one sentence.", 2)
sep()
ans_header()
ans("Engineering: applying scientific knowledge to build cost-effective solutions to practical problems. [1]")
ans("Software Engineering: the design, development, testing, and maintenance of software applications, applying engineering principles and programming knowledge to build solutions for end users. [1]")

q("List three reasons why software projects fail.", 3)
sep()
ans_header()
ans("Poor requirements (incomplete, over-ambitious, or unnecessary). [1]")
ans("Over budget / poor estimation. [1]")
ans("Contract management issues / end-user training gaps / operational management failures. [1]")
ans_text("Also acceptable: bad communication, scope creep, lack of stakeholder involvement.", italic=True)

q("List the five stages of the Waterfall SDLC.", 5)
sep()
ans_header()
ans("Requirements Definition [1]")
ans("System and Software Design [1]")
ans("Implementation and Unit Testing [1]")
ans("Integration and System Testing [1]")
ans("Operation and Maintenance [1]")

q("Explain the difference between Verification and Validation.", 3)
sep()
ans_header()
ans("Verification: \"Are we building it right?\" — checks software complies with specifications, constraints, and regulations. [1]")
ans("Validation: \"Are we building the right thing?\" — checks software meets the actual needs of customers/stakeholders. [1]")
ans("A system can pass verification but fail validation if the specification itself does not address user needs. [1]")

q("State two advantages and two disadvantages of the Waterfall model.", 2)
sep()
ans_header()
ans_text("Advantages (0.5 each):", bold=True)
ans("Simple to use and understand; every phase has a defined result and review.")
ans_text("Disadvantages (0.5 each):", bold=True)
ans("Software is ready only after the last stage; not suited for long-term projects where requirements change.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# L2: AGILE (22 marks)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 2 — Agile Software Development  (22 marks)", level=1)

q("State the four key values of the Agile Manifesto.", 4)
sep()
ans_header()
ans("Individuals and interactions over processes and tools. [1]")
ans("Working software over comprehensive documentation. [1]")
ans("Customer collaboration over contract negotiation. [1]")
ans("Responding to change over following a plan. [1]")

q("List the five elements of the Extreme Programming (XP) ethos.", 5)
sep()
ans_header()
ans("Simple design: use the simplest way to implement features. [1]")
ans("Sustainable pace: effort is constant and manageable. [1]")
ans("Coding standards: teams follow an agreed style and format. [1]")
ans("Collective ownership: everyone owns all the code. [1]")
ans("Whole team approach: everyone is included in everything. [1]")

q("List the five Extreme Programming (XP) practices.", 5)
sep()
ans_header()
ans("Pair programming [1]")
ans("Test-driven development [1]")
ans("Small releases [1]")
ans("Continuous integration [1]")
ans("Refactoring [1]")

q("In pair programming, what are the two roles and their responsibilities?", 2)
sep()
ans_header()
ans("Helm: uses the keyboard and mouse, does the actual coding. [1]")
ans("Tactician: thinks about implications, potential problems, and is positioned to recommend refactoring. [1]")

q("State three benefits of Test-Driven Development (TDD).", 3)
sep()
ans_header()
ans("Code coverage: all code has at least one test because code isn't written without a test. [1]")
ans("Simplified debugging: if a test fails, it was caused by the last change. [1]")
ans("System documentation: tests describe what the code should be doing. [1]")

q("List three problems or limitations of Agile development.", 3)
sep()
ans_header()
ans("Hard to draw up legally binding contracts — no full spec in advance. [1]")
ans("Good for greenfield, not as effective for brownfield/legacy systems. [1]")
ans("Works for small co-located teams, challenging for large distributed development. [1]")
ans_text("Also acceptable: reliance on developer knowledge (holiday/illness/turnover).", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# L3: REQUIREMENTS (23 marks)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 3 — Requirements Engineering  (23 marks)", level=1)

q("Define functional requirements and non-functional requirements.", 2)
sep()
ans_header()
ans("Functional: specify user interactions — what the system should do (services, reactions to inputs, behaviors). [1]")
ans("Non-functional: specify quality properties — how functional requirements are realized (constraints, security, usability, performance). [1]")

q("What does INVEST stand for? Explain each letter.", 6)
sep()
ans_header()
ans("Independent: can be worked on separately from other stories. [1]")
ans("Negotiable: flexible and open to discussion. [1]")
ans("Valuable: delivers clear value to the user. [1]")
ans("Estimable: can be estimated for effort. [1]")
ans("Small: small enough to complete within a sprint. [1]")
ans("Testable: has clear criteria to determine if it's done. [1]")

q("Write the template for a User Story.", 1)
sep()
ans_header()
ans("As a <type of user>, I want to <some goal> so that <some reason>. [1]")

q("Write the template for Acceptance Criteria.", 1)
sep()
ans_header()
ans("Given <initial context/precondition>, When <action/event>, Then <expected outcome>. [1]")

q("State the four qualities that good acceptance criteria must have.", 2)
sep()
ans_header()
ans("Clear (unambiguous), Testable (can verify), Measurable (quantitative/qualitative), Atomic (independent, checked by itself). [0.5 each]")

q("Explain the hierarchy: Initiative → Epic → User Story.", 3)
sep()
ans_header()
ans("Initiative: strategic objective with important business outcome, spans multiple epics and teams. [1]")
ans("Epic: large strategic goal, spans multiple sprints, needs to be broken down. [1]")
ans("User Story: specific feature/functionality, completed within a single sprint. [1]")

q("Explain the MoSCoW prioritization method.", 2)
sep()
ans_header()
ans("Must-Have: essential. Should-Have: important. Could-Have: nice to have. Won't-Have: out of scope at present. [0.5 each]")

q("What is a Use Case? What is an Actor?", 3)
sep()
ans_header()
ans("Use Case: a set of use-case instances — each a sequence of actions the system performs that yields an observable result of value to a particular actor. Models a dialogue between actors and the system. [1.5]")
ans("Actor: represents anything that interacts with the system (human, machine, or another system). Actors are EXTERNAL — not part of the system. [1.5]")

q("What is the Onion Model used for in requirements engineering?", 1)
sep()
ans_header()
ans("It is a template for identifying stakeholders — organising them in concentric layers from the product/service outward to the wider environment (operators, sponsors, developers, regulators, public, etc.). [1]")

q("What are the two parts of a use-case model?", 2)
sep()
ans_header()
ans("Use-case diagrams (visual representation showing actors, use cases, and relationships). [1]")
ans("Use-case specifications (text: description of flows of events, preconditions, postconditions, key scenarios). [1]")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# L4: OOD (25 marks)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 4 — Object-Oriented Design  (25 marks)", level=1)

q("Define Encapsulation and explain why it matters.", 2)
sep()
ans_header()
ans("Encapsulation: encloses the data and behaviour of an object within that object. [1]")
ans("It prevents unauthorised access and protects data integrity — objects can only be changed through permitted (public) methods. [1]")

q("Define Abstraction and explain its benefit.", 2)
sep()
ans_header()
ans("Abstraction: focuses on core concerns, exposes only essential information, and hides internal complexity. [1]")
ans("Benefit: promotes loose coupling — objects interact via abstract interfaces rather than concrete implementations. [1]")

q("Define Inheritance and state two benefits.", 2)
sep()
ans_header()
ans("Inheritance: a class inherits the properties and behaviour of another class (\"is-a\" relationship). [1]")
ans("Benefits: reuse code (reduce re-writing), reduce errors and inconsistency. [1]")

q("Define Polymorphism and give a brief example.", 2)
sep()
ans_header()
ans("Polymorphism: allows a subclass to be substituted for its superclass; behaviour belongs to the subclass object (dynamic method dispatch). [1]")
ans("Example: Dog and Cat both extend Animal. Calling makeSound() on an Animal reference invokes the specific subclass implementation (Woof vs Meow). [1]")

q("Define Composition and distinguish it from Aggregation.", 3)
sep()
ans_header()
ans("Composition: builds an object from other objects with strong ownership and coincident lifetimes — parts cannot survive without the whole. Non-shared aggregation. [1.5]")
ans("Aggregation: a whole-part ('is a part-of') relationship where parts CAN exist independently. Can be shared (part belongs to multiple wholes). [1.5]")

q("What does a UML Class Diagram represent?", 1)
sep()
ans_header()
ans("A static view of a system — shows classes (name, attributes, methods with visibility), and the structural relationships between them. [1]")

q("List the four visibility/access modifiers in UML and their meanings.", 2)
sep()
ans_header()
ans("+ public: available to all. [0.5]")
ans("- private: available only within the class. [0.5]")
ans("# protected: available to subclasses and same-package classes. [0.5]")
ans("~ package: available only within the package. [0.5]")

q("Define Association and Multiplicity in class diagrams.", 2)
sep()
ans_header()
ans("Association: a structural relationship specifying that objects of one class are connected to objects of another. [1]")
ans("Multiplicity: the number of instances one class relates to ONE instance of another class. E.g., 1, 0..*, 1..*, 0..1. [1]")

q("Explain the difference between Generalization and Aggregation.", 2)
sep()
ans_header()
ans("Generalization: an 'is a kind of' relationship — subclass inherits from superclass (e.g., NHS Patient is-a Patient). [1]")
ans("Aggregation: a 'part-of' relationship — one object contains others as parts (e.g., Hospital has-a Department). [1]")

q("What is the difference between Abstract and Concrete classes?", 2)
sep()
ans_header()
ans("Abstract classes cannot be instantiated (no direct objects). They serve as blueprints for subclasses. [1]")
ans("Concrete classes can be instantiated — they have direct objects. [1]")

q("What is the purpose of a Sequence Diagram? List four key elements.", 3)
sep()
ans_header()
ans("Purpose: models how objects collaborate and interact over time through messages. Good for real-time specs and complex scenarios. [1]")
ans("Participants (actors and objects), Lifeline (vertical dashed line = existence over time), Messages (horizontal arrows = requests between objects), Return messages (dashed arrows = values returned). [0.5 each = 2]")

q("What interaction frames can be used in sequence diagrams?", 2)
sep()
ans_header()
ans("alt — branching / conditional logic (like if-else). [1]")
ans("loop — iteration / repeating a sequence. [1]")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# L5: PROJECT MANAGEMENT (15 marks)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Part 5 — Project Management  (15 marks)", level=1)

q("What are Story Points and why do we use them?", 2)
sep()
ans_header()
ans("Story Points: an informal agile unit of 'size measurement' (usually 1–10), used to estimate relative effort for user stories. [1]")
ans("Based on the 'Wisdom of the Crowds' — collective team estimates are more accurate than individual estimates. [1]")

q("Describe how Planning Poker works.", 3)
sep()
ans_header()
ans("Each team member receives numbered cards following the Fibonacci sequence (1, 3, 5, 8, 13, 20...). [1]")
ans("For each story, all members simultaneously reveal their estimate card. Diverging estimates are discussed. [1]")
ans("Process repeats for up to 3 iterations. Fibonacci spacing reflects that larger tasks are harder to estimate precisely. [1]")

q("Define Team Velocity and explain how it helps planning.", 2)
sep()
ans_header()
ans("Team Velocity = number of estimated story points completed per sprint, derived from previous sprints (e.g., average). [1]")
ans("Used to estimate time to complete remaining backlog and to set sprint targets. [1]")

q("Compare white-box and black-box complexity metrics.", 2)
sep()
ans_header()
ans("White-box (after development, based on source code): e.g., Lines of Code, Cyclomatic Complexity V(G)=E−N+2P. [1]")
ans("Black-box (before development, based on requirements/specs): e.g., Story Points, Planning Poker. Used when no code exists yet. [1]")

q("Describe the key characteristics of Sprint Review (Purpose, Attendees, Output, Key Question).", 4)
sep()
ans_header()
ans("Purpose: inspect the product increment and gather stakeholder feedback. [1]")
ans("Attendees: dev team, Product Owner, Scrum Master, stakeholders, customers, users. [1]")
ans("Output: feedback on product, potential changes to backlog/requirements. [1]")
ans("Key Question: 'Does the product meet stakeholder expectations, and what adjustments are needed?' [1]")

q("Describe the key characteristics of Sprint Retrospective (Purpose, Attendees, Output, Key Question).", 2)
sep()
ans_header()
ans("Purpose: inspect the process, reflect on team performance, identify improvements. Attendees: dev team + Scrum Master (PO optional). [1]")
ans("Output: actionable items for improving team processes. Key Question: 'How can we improve the way we work as a team?' [1]")

# ═══ End ═══
doc.add_page_break()
doc.add_heading("End of Quiz", level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
total = q_num[0]
p.add_run(f"Total: {total} questions · 100 marks").bold = True

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SEDP_Concept_Quiz_L1-L5.docx")
doc.save(out_path)
print(f"Generated {total} questions → {out_path}")
