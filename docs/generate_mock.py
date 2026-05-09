"""Generate a mock exam Word document for SEDP L1-L5."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ─── Styles ──────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_q(text, level=2):
    doc.add_heading(text, level=level)

def add_p(text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Inches(0.3 + indent * 0.3)
    p.text = ""
    p.add_run(text)
    return p

def add_answer_header():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Sample Answer")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
    run.font.size = Pt(11)

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("COMSM0166")
r.font.size = Pt(16)
r.bold = True

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Software Engineering Discipline and Practice")
r2.font.size = Pt(14)

doc.add_paragraph()

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Practice Paper — Chapters 1–5")
r3.font.size = Pt(18)
r3.bold = True
r3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

for line in [
    "TIME ALLOWED: 120 minutes (2 hours)",
    "",
    "Answer ALL questions in Section A and TWO questions in Section B.",
    "",
    "Maximum marks: 80",
    "",
    "This is a closed book exam.",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    if "80" in line or "ALL" in line:
        run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION A
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("SECTION A: Answer ALL questions (30 marks)", level=1)
doc.add_paragraph()

# ─── Question 1 ───────────────────────────────────────────────────────
add_q("Question 1.  [10 marks]  —  Software Development Lifecycle")

add_p("a) List and briefly describe the five stages of the Waterfall SDLC. [5 marks]")
add_p("b) Explain the difference between Verification and Validation. Give an example of a scenario where software passes verification but fails validation. [5 marks]")

add_separator()
add_answer_header()

add_p("a) Five stages of Waterfall (1 mark each):", bold=True)
add_bullet("Requirements Definition — gathering and documenting what the system should do (functional and non-functional requirements).")
add_bullet("System and Software Design — deciding how parts of the system interact, dividing work, API partitioning.")
add_bullet("Implementation and Unit Testing — parallel coding, version control, automated build, writing documentation.")
add_bullet("Integration and System Testing — combining modules and verifying the whole system works together.")
add_bullet("Operation and Maintenance — deploying the system, ongoing support, bug fixes, and updates.")

add_p("b) Verification vs Validation (5 marks):", bold=True)
add_bullet("Verification: \"Are we building it right?\" — checks the software complies with its specifications, constraints, and regulations. [1 mark]")
add_bullet("Validation: \"Are we building the right thing?\" — checks the software meets the actual needs of customers and stakeholders. [1 mark]")
add_bullet("Key difference: verification is about conformance to spec; validation is about fitness for purpose. [1 mark]")
add_bullet("Example: A team builds a payroll system exactly matching the written specification (passes verification), but the specification failed to capture that employees needed to view payslips on mobile devices. The system fails validation because it doesn't meet actual user needs. [2 marks]")

doc.add_page_break()

# ─── Question 2 ───────────────────────────────────────────────────────
add_q("Question 2.  [10 marks]  —  Agile Fundamentals")

add_p("a) State the four key values of the Agile Manifesto. [4 marks]")
add_p("b) Describe three problems or limitations of Agile software development. [3 marks]")
add_p("c) Explain two key differences between Test-Driven Development (TDD) and traditional testing. State two benefits of TDD. [3 marks]")

add_separator()
add_answer_header()

add_p("a) Four Agile values (1 mark each):", bold=True)
add_bullet("Individuals and interactions over processes and tools.")
add_bullet("Working software over comprehensive documentation.")
add_bullet("Customer collaboration over contract negotiation.")
add_bullet("Responding to change over following a plan.")

add_p("b) Three problems with Agile (1 mark each):", bold=True)
add_bullet("Hard to draw up legally binding contracts because a full specification is never written in advance.")
add_bullet("Good for greenfield development but less effective for brownfield/legacy systems.")
add_bullet("Works well for small co-located teams but challenging for large distributed development.")
add_p("Also acceptable: reliance on developer knowledge (holiday/illness/turnover risk).", italic=True)

add_p("c) TDD vs traditional testing (3 marks):", bold=True)
add_bullet("In TDD, tests are written before any code; in traditional testing, tests are written after code is implemented. [1 mark]")
add_bullet("In TDD, if there is no test for a feature, the feature is not implemented; code is driven by tests. [0.5 marks]")
add_bullet("Benefits (0.5 marks each): (1) Code coverage — all written code has at least one test. (2) Simplified debugging — a failing test must be caused by the last change. Also acceptable: tests serve as system documentation.")

doc.add_page_break()

# ─── Question 3 ───────────────────────────────────────────────────────
add_q("Question 3.  [10 marks]  —  Estimation and Measurement")

add_p("a) Explain what Story Points are and how Planning Poker works. [4 marks]")
add_p("b) Define Team Velocity and explain how it is used to estimate project timelines. [3 marks]")
add_p("c) Compare Lines of Code (LoC) and Cyclomatic Complexity as white-box metrics. State one advantage and one limitation of each. [3 marks]")

add_separator()
add_answer_header()

add_p("a) Story Points and Planning Poker (4 marks):", bold=True)
add_bullet("Story Points are an informal agile unit of size estimation, usually 1–10, representing relative effort for a user story. [1 mark]")
add_bullet("Estimates are derived collectively by the whole team, based on the 'Wisdom of the Crowds' principle. [1 mark]")
add_bullet("Planning Poker: each team member has cards following the Fibonacci sequence (1, 3, 5, 8, 13, 20...). [0.5 marks]")
add_bullet("All members simultaneously reveal their card for a story. If estimates diverge, the team discusses and re-votes, up to 3 iterations. [0.5 marks]")
add_bullet("Fibonacci spacing reflects that larger tasks are harder to estimate precisely. Low = trivial, High = difficult. [1 mark]")

add_p("b) Team Velocity (3 marks):", bold=True)
add_bullet("Team Velocity = the number of estimated story points completed per sprint. [1 mark]")
add_bullet("Derived from previous sprints (e.g., average of last N sprints). [1 mark]")
add_bullet("Used to estimate: (1) time required to complete remaining backlog, (2) target number of stories for the next sprint. [1 mark]")

add_p("c) LoC vs Cyclomatic Complexity (3 marks):", bold=True)
add_bullet("Lines of Code (LoC): Advantage — easy to compute and widely used. Limitation — ignores logical/architectural complexity, highly language-specific, not all lines are equal. [1.5 marks]")
add_bullet("Cyclomatic Complexity V(G) = E − N + 2P: Advantage — measures the number of independent paths, giving a better picture of logical complexity. Limitation — does not capture code size, requires building a control flow graph. [1.5 marks]")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION B
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("SECTION B: Answer TWO of THREE questions (50 marks)", level=1)
doc.add_paragraph()

# ─── Question 4 ───────────────────────────────────────────────────────
add_q("Question 4.  [25 marks]  —  Requirements Engineering")

add_p("You are part of an agile team building a Student Accommodation Portal for a university. The system helps students search for accommodation, submit applications, and communicate with landlords. Landlords can list properties, manage tenants, and track payments.", bold=False, italic=True)
doc.add_paragraph()

add_p("4.1  Identify three types of stakeholders using the Onion Model and describe their relationship to the system. [3 marks]")
add_p("4.2  Write 2 epics and 3 user stories per epic for this system. [12 marks]")
add_p("4.3  Write acceptance criteria (Given/When/Then format) for two of the user stories above. [4 marks]")
add_p("4.4  Evaluate this user story using INVEST criteria: \"As a user, I want to manage my accommodation so that everything works well.\" Identify which criteria it fails and explain why. [3 marks]")
add_p("4.5  Explain the difference between functional and non-functional requirements, giving one example of each for this system. [3 marks]")

add_separator()
add_answer_header()

add_p("4.1  Stakeholders (1 mark each):", bold=True)
add_bullet("Students (Functional Beneficiary / Normal Operator) — primary users who search and apply for accommodation.")
add_bullet("Landlords (Functional Beneficiary) — list properties, manage tenants, interact with the system directly.")
add_bullet("University Administration (Sponsor / Champion) — funds and promotes the system, sets policies.")
add_p("Also acceptable: IT support (Maintenance Operator), Regulators (external), Negative stakeholders (competing platforms).", italic=True)

add_p("4.2  Epics and User Stories (2 marks per epic title + 1 mark per story = 12 marks):", bold=True)
add_p("Epic 1: Accommodation Search and Discovery", bold=True, indent=1)
add_bullet("As a student, I want to search for accommodation by location, price, and room type so that I can find suitable options quickly.", indent=1)
add_bullet("As a student, I want to view photos and descriptions of a property so that I can assess it before visiting.", indent=1)
add_bullet("As a student, I want to save properties to a shortlist so that I can compare them later.", indent=1)
add_p("Epic 2: Application and Booking", bold=True, indent=1)
add_bullet("As a student, I want to submit an application for a property so that the landlord can review my request.", indent=1)
add_bullet("As a landlord, I want to review and approve/reject student applications so that I can select suitable tenants.", indent=1)
add_bullet("As a student, I want to receive a notification when my application status changes so that I know the outcome promptly.", indent=1)

add_p("4.3  Acceptance Criteria (2 marks each, 2 stories):", bold=True)
add_p("Story: \"Search for accommodation by location, price, and room type\"", italic=True, indent=1)
add_bullet("Given the student is on the search page, When they enter a location, set a price range, and select a room type, Then the system displays a filtered list of matching properties.", indent=1)
add_p("Story: \"Submit an application for a property\"", italic=True, indent=1)
add_bullet("Given the student is viewing a property detail page, When they click 'Apply' and confirm their details, Then the system records the application and notifies the landlord via email.", indent=1)

add_p("4.4  INVEST evaluation (1 mark per failed criterion, 3 marks):", bold=True)
add_p("\"As a user, I want to manage my accommodation so that everything works well.\"", italic=True, indent=1)
add_bullet("Fails Small — 'manage my accommodation' is an initiative/epic, not a sprint-sized story. [1 mark]", indent=1)
add_bullet("Fails Estimable — too vague to estimate effort; unclear what 'manage' entails. [1 mark]", indent=1)
add_bullet("Fails Testable — 'everything works well' has no clear acceptance criteria. [1 mark]", indent=1)

add_p("4.5  Functional vs Non-functional (3 marks):", bold=True)
add_bullet("Functional: specifies what the system should do — e.g., \"The system shall allow students to filter properties by price range.\" [1.5 marks]")
add_bullet("Non-functional: specifies quality properties of operation — e.g., \"The search results page shall load within 2 seconds under normal load.\" [1.5 marks]")

doc.add_page_break()

# ─── Question 5 ───────────────────────────────────────────────────────
add_q("Question 5.  [25 marks]  —  Object-Oriented Design")

add_p("You are designing an Online Food Delivery System. Customers browse restaurant menus, place orders, and pay online. Each restaurant manages its own menu of food items. A delivery driver is assigned to each order. The system tracks order status from placement through delivery.", bold=False, italic=True)
doc.add_paragraph()

add_p("5.1  Sketch a UML class diagram with at least 5 main classes, showing attributes, methods, relationships, and multiplicities. [7 marks]")
add_p("5.2  List 2 key attributes and 2 key methods for the Order class. Explain the access modifiers you would use and why. [4 marks]")
add_p("5.3  Using examples from your design, provide a one-sentence explanation for each of the five OO principles: Encapsulation, Abstraction, Inheritance, Polymorphism, and Composition. [10 marks]")
add_p("5.4  Draw a sequence diagram for the 'Place Order' scenario, showing interactions between Customer, Order, Restaurant, and DeliveryDriver. [4 marks]")

add_separator()
add_answer_header()

add_p("5.1  Class Diagram (7 marks):", bold=True)
add_p("Expected classes (1 mark per class, up to 5):", indent=1)
add_bullet("Customer — name, address, email; placeOrder(), viewHistory()", indent=1)
add_bullet("Restaurant — name, location; addMenuItem(), updateMenu()", indent=1)
add_bullet("MenuItem — name, price, category; updatePrice()", indent=1)
add_bullet("Order — orderId, totalAmount, status, timestamp; calculateTotal(), updateStatus()", indent=1)
add_bullet("DeliveryDriver — name, vehicleType, isAvailable; acceptDelivery(), updateLocation()", indent=1)
add_p("Relationships shown [1 mark]: Customer → Order (1 to *), Order → Restaurant (many to 1), Restaurant → MenuItem (1 to *), Order → DeliveryDriver (1 to 0..1).", indent=1)
add_p("Multiplicities labelled [1 mark].", indent=1)

add_p("5.2  Order class details (4 marks):", bold=True)
add_p("Attributes (1 mark each):", indent=1)
add_bullet("- orderId : String (private) — unique identifier, should not be modified externally.", indent=1)
add_bullet("- totalAmount : Double (private) — sensitive financial data, protected from direct manipulation.", indent=1)
add_p("Methods (1 mark each):", indent=1)
add_bullet("+ calculateTotal() : Double (public) — needs to be called by other classes to compute the bill.", indent=1)
add_bullet("+ updateStatus(newStatus : String) (public) — allows Restaurant and Driver to update order progress.", indent=1)

add_p("5.3  OO Principles (2 marks each = 10 marks):", bold=True)
add_bullet("Encapsulation: Order has private attributes (totalAmount, orderId) with public methods (calculateTotal, updateStatus). This protects data integrity by preventing direct external modification. [2 marks]")
add_bullet("Abstraction: A Payment class could be abstract, with subclasses CreditCardPayment and CashPayment. This hides payment processing complexity from the rest of the system. [2 marks]")
add_bullet("Inheritance: PremiumCustomer and RegularCustomer inherit from Customer, reusing common attributes (name, address) while adding specialised behaviour (discount calculation). [2 marks]")
add_bullet("Polymorphism: calculateDeliveryFee() behaves differently in BikeDriver and CarDriver subclasses. The system calls the method on a DeliveryDriver reference without knowing the specific subclass. [2 marks]")
add_bullet("Composition: An Order is composed of OrderItems (which reference MenuItems). If the Order is deleted, its OrderItems cease to exist — they share the same lifetime. [2 marks]")

add_p("5.4  Sequence Diagram (4 marks):", bold=True)
add_bullet("Participants: Customer, :Order, :Restaurant, :DeliveryDriver [1 mark]")
add_bullet("Customer → Order: createOrder(items) [0.5 marks]")
add_bullet("Order → Restaurant: confirmItems(items) [0.5 marks]")
add_bullet("Restaurant → Order: return confirmation + estimated time [0.5 marks]")
add_bullet("Order → Order: calculateTotal() [0.5 marks]")
add_bullet("Order → DeliveryDriver: assignDriver() [0.5 marks]")
add_bullet("Order → Customer: return orderConfirmation (orderId, total, ETA) [0.5 marks]")

doc.add_page_break()

# ─── Question 6 ───────────────────────────────────────────────────────
add_q("Question 6.  [25 marks]  —  Scrum and Project Management")

add_p("FitTrack is a startup building a fitness tracking mobile app. The team uses Scrum with 2-week sprints. After 4 sprints, the completed story points are: Sprint 1 = 18, Sprint 2 = 22, Sprint 3 = 20, Sprint 4 = 24. The remaining product backlog contains 126 story points.", bold=False, italic=True)
doc.add_paragraph()

add_p("6.1  Describe the three key roles in the Scrum framework and their responsibilities. [6 marks]")
add_p("6.2  Compare Sprint Review and Sprint Retrospective by completing a table with the following rows: Purpose, Attendees, Output, What Is Discussed, Key Question. [10 marks]")
add_p("6.3  Using the velocity data above, calculate the team's average velocity and estimate how many more sprints are needed to complete the remaining backlog. [4 marks]")
add_p("6.4  Describe what a Burn-Down Chart shows and explain how the feedback loop between estimated and actual progress helps the team. [5 marks]")

add_separator()
add_answer_header()

add_p("6.1  Scrum Roles (2 marks each):", bold=True)
add_bullet("Product Owner: sets the product vision, prioritises the product backlog, writes epics and user stories, ensures alignment with business needs. Acts as the voice of the customer. [2 marks]")
add_bullet("Scrum Master: facilitates agile processes (daily stand-ups, sprint planning, reviews, retrospectives), removes obstacles, ensures the team follows agile principles. [2 marks]")
add_bullet("Development Team: self-organising group that writes, tests, and deploys code. Responsible for delivering the sprint backlog as working software. [2 marks]")

add_p("6.2  Sprint Review vs Retrospective (1 mark per cell = 10 marks):", bold=True)

table = doc.add_table(rows=6, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["", "Sprint Review", "Sprint Retrospective"]
data = [
    ["Purpose", "Inspect the product increment and gather stakeholder feedback.", "Inspect the process, reflect on team performance, identify improvements."],
    ["Attendees", "Dev team, Product Owner, Scrum Master, stakeholders, customers.", "Dev team and Scrum Master (PO optional)."],
    ["Output", "Feedback on product; potential backlog/requirement changes.", "Actionable items to improve team processes in future sprints."],
    ["What Is\nDiscussed", "What was completed; demo of increment; stakeholder feedback; priority changes.", "What went well; what didn't; blockers; how to improve collaboration."],
    ["Key Question", "\"Does the product meet stakeholder expectations?\"", "\"How can we improve the way we work as a team?\""],
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_text

# Set font size for table
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)

doc.add_paragraph()

add_p("6.3  Velocity Calculation (4 marks):", bold=True)
add_bullet("Average velocity = (18 + 22 + 20 + 24) / 4 = 84 / 4 = 21 story points per sprint. [2 marks]")
add_bullet("Remaining backlog = 126 points. Sprints needed = 126 / 21 = 6 more sprints. [1 mark]")
add_bullet("At 2 weeks per sprint, this means approximately 12 more weeks. [1 mark]")

add_p("6.4  Burn-Down Chart and Feedback Loops (5 marks):", bold=True)
add_bullet("A burn-down chart plots remaining work (story points) on the Y-axis against time/sprints on the X-axis. [1 mark]")
add_bullet("It shows two lines: the ideal/estimated line (straight diagonal from total to zero) and the actual progress line. [1 mark]")
add_bullet("If actual is above ideal, the team is behind schedule; if below, they are ahead. [1 mark]")
add_bullet("The feedback loop: after each sprint, the team compares actual progress against the estimate. Discrepancies trigger discussion in the Sprint Retrospective about causes and corrective actions. [1 mark]")
add_bullet("This enables continuous improvement: the team can adjust velocity estimates, reprioritise backlog items, or identify and remove blockers for future sprints. [1 mark]")


# ═══════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, "SEDP_Practice_Paper_L1-L5.docx")
doc.save(out_path)
print(f"Saved to {out_path}")
